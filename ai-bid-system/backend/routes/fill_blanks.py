import io
import json
from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from services.text_replacer import replace_text_fuzzy
from routes.merge_docs import fill_dynamic_tables, find_and_replace_insert_codes

router = APIRouter()


@router.post("/fill-blanks")
async def fill_blanks(
    file: UploadFile = File(...),
    normal_blanks: str = Form("[]"),
    dynamic_tables: str = Form("[]"),
    mapping: str = Form("{}")
):
    """
    填充Word文档的所有空白（普通填空 + 动态表格 + 附件合并）
    
    Args:
        file: 上传的Word模板文件
        normal_blanks: JSON字符串，格式：[{"paraIndex": 0, "originalText": "___", "filledText": "张三"}]
        dynamic_tables: JSON字符串，动态表格数据（复用现有格式）
        mapping: JSON字符串，附件URL映射（复用现有格式）
    
    Returns:
        StreamingResponse: 填充后的Word文档
    """
    print("\n" + "="*60)
    print("📥 文档填充请求 (/fill-blanks)")
    print("="*60)
    
    try:
        # 1. 读取上传的文档
        doc_bytes = await file.read()
        print(f"📥 主文档: {len(doc_bytes)} bytes")
        
        main_doc = Document(io.BytesIO(doc_bytes))
        print(f"📄 段落数: {len(main_doc.paragraphs)}, 表格数: {len(main_doc.tables)}")
        
        # ✅ 使用 xpath 获取所有 <w:p> 节点（含表格内的段落），与扫描时保持一致
        all_p_elements = main_doc._element.xpath('//w:p')
        all_paragraphs = [Paragraph(p, main_doc._part) for p in all_p_elements]
        print(f"📄 全局段落数(含表格内): {len(all_paragraphs)}, 主段落(body only): {len(main_doc.paragraphs)}")
        
        # 2. 解析 normal_blanks
        parsed_blanks = []
        if normal_blanks:
            try:
                parsed_blanks = json.loads(normal_blanks)
                print(f"📝 普通填空数据: {len(parsed_blanks)} 个")
            except json.JSONDecodeError as e:
                print(f"⚠️ normal_blanks JSON解析失败: {e}")
        
        # 3. 处理普通填空（核心逻辑）
        if parsed_blanks:
            print(f"\n{'='*60}")
            print("🔧 开始填充普通空白")
            print(f"{'='*60}")
            
            filled_count = 0
            for blank in parsed_blanks:
                para_index = blank.get("paraIndex")
                original_text = blank.get("originalText", "")
                filled_text = blank.get("filledText", "")
                blank_type = blank.get("type", "")
                blank_context = blank.get("context", "")
                
                # 验证参数
                if para_index is None or not filled_text:
                    print(f"⚠️ 跳过无效填空: paraIndex={para_index}, filledText={filled_text}")
                    continue
                
                # 验证 paraIndex 范围（使用全局段落列表）
                if para_index < 0 or para_index >= len(all_paragraphs):
                    print(f"❌ paraIndex {para_index} 超出范围 (文档共 {len(all_paragraphs)} 个全局段落)")
                    continue
                
                # 获取目标段落（使用全局 xpath 索引，与扫描时一致）
                target_para = all_paragraphs[para_index]
                
                # 执行替换（使用模糊匹配策略，支持图片URL自动插入）
                success = replace_text_fuzzy(target_para, original_text, filled_text, blank_type, blank_context)
                
                if success:
                    filled_count += 1
                    print(f"   ✅ [{para_index}] '{original_text[:30]}...' → '{filled_text[:30]}...'")
                else:
                    print(f"   ⚠️ [{para_index}] 替换失败: '{original_text[:30]}...'")
                    print(f"      段落内容: {target_para.text[:100]}...")
            
            print(f"✅ 普通填空完成，成功填充 {filled_count}/{len(parsed_blanks)} 个")
        
        # 4. 处理动态表格（复用现有逻辑）
        parsed_dynamic = []
        if dynamic_tables:
            try:
                parsed_dynamic = json.loads(dynamic_tables)
                print(f"📋 动态表格数据: {len(parsed_dynamic)} 个表格")
            except json.JSONDecodeError as e:
                print(f"⚠️ dynamic_tables JSON解析失败: {e}")
        
        if parsed_dynamic:
            print(f"\n{'='*60}")
            print("🔧 开始填充动态表格")
            print(f"{'='*60}")
            filled = fill_dynamic_tables(main_doc, parsed_dynamic)
            print(f"✅ 动态表格填充完成，共 {filled} 行")
        
        # 5. 处理附件合并（复用现有逻辑）
        url_mapping = {}
        if mapping:
            try:
                url_mapping = json.loads(mapping)
            except json.JSONDecodeError:
                url_mapping = {}
        
        if url_mapping:
            print(f"\n{'='*60}")
            print("🔍 扫描 [INSERT_DOC:xxx] 暗码并合并子文档")
            print(f"{'='*60}")
            print(f"📋 mapping: {len(url_mapping)} 条")
            total_merged = find_and_replace_insert_codes(main_doc, url_mapping)
            print(f"📦 合并完成，共 {total_merged} 个附件")
        
        # 6. 保存并返回文档
        output_io = io.BytesIO()
        main_doc.save(output_io)
        output_io.seek(0)
        
        final_size = len(output_io.getvalue())
        print(f"✅ 最终文档大小: {final_size} bytes\n")
        
        return StreamingResponse(
            output_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Filled_Document.docx"}
        )
    
    except Exception as e:
        print(f"💥 服务器错误: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": str(e), "type": type(e).__name__}
        )
