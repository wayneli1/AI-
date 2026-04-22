import copy
import io
import json
import re

import requests
import urllib3
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Cm
from docxcompose.composer import Composer
from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

router = APIRouter()


def normalize_key(key):
    if not key:
        return key
    return re.sub(r'\s+', '', key)


def find_url_fuzzy(code, url_mapping):
    if code in url_mapping:
        return url_mapping[code]

    normalized_code = normalize_key(code)
    for key, url in url_mapping.items():
        if normalize_key(key) == normalized_code:
            return url

    full_code = f"[INSERT_DOC:{code}]"
    if full_code in url_mapping:
        return url_mapping[full_code]

    normalized_full = normalize_key(full_code)
    for key, url in url_mapping.items():
        if normalize_key(key) == normalized_full:
            return url

    return None


def _build_heading_level_map(doc):
    """
    根据文档的 styles 表动态构建样式 ID -> 标题级别的映射。
    这样无论文档使用什么样式 ID 约定，都能正确识别标题。
    """
    mapping = {}
    try:
        for style in doc.styles:
            if style.name:
                m = re.match(r'^Heading\s*(\d+)$', style.name, re.IGNORECASE)
                if m:
                    mapping[style.style_id] = int(m.group(1))
    except Exception as e:
        print(f"   [WARN] 构建标题映射失败: {e}")
    return mapping


def shift_heading_levels(element, doc):
    heading_level_map = _build_heading_level_map(doc)

    qn_pStyle = qn('w:pStyle')
    qn_numPr = qn('w:numPr')

    for p in element.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue

        # 🌟 提取段落纯文本内容
        text = "".join(t.text for t in p.iter(qn('w:t')) if t.text).strip()

        style_elem = pPr.find(qn_pStyle)
        is_heading = False
        is_body_text_2 = False

        if style_elem is not None:
            val = style_elem.get(qn('w:val'), '')

            # 判断是否为标题
            m = re.match(r'^Heading\s*(\d+)$', val, re.IGNORECASE)
            is_heading_match = bool(m or val in heading_level_map)
            
            # 判断是否为“正文文本”的变体（罪魁祸首：ID为 2 或 3）
            is_body_text_variant = val in ['2', '3'] or 'bodytext' in val.lower().replace(' ', '')

            if is_heading_match or is_body_text_variant:
                # 🛡️ 绝杀拦截：如果是长文本或以冒号/句号结尾，统统打回原形
                if len(text) > 35 or text.endswith(('。', '.', '：', ':')) or not text:
                    style_elem.set(qn('w:val'), 'Normal')
                    is_body_text_2 = True
                else:
                    if is_heading_match:
                        is_heading = True
                        if m:
                            old_level = int(m.group(1))
                            new_level = min(old_level + 1, 9)
                            style_elem.set(qn('w:val'), f'Heading{new_level}')
                        elif val in heading_level_map:
                            old_level = heading_level_map[val]
                            new_level = min(old_level + 1, 9)
                            new_style_id = None
                            for sid, level in heading_level_map.items():
                                if level == new_level:
                                    new_style_id = sid
                                    break
                            if new_style_id:
                                style_elem.set(qn('w:val'), new_style_id)
                            else:
                                style_elem.set(qn('w:val'), f'Heading{new_level}')
                    elif is_body_text_variant:
                        # 确实是短文本的 Body Text，为了防止 ID 冲突也洗成 Normal
                        style_elem.set(qn('w:val'), 'Normal')
                        is_body_text_2 = True

        # 剥离可能残留的编号属性
        if is_heading or is_body_text_2:
            numPr = pPr.find(qn_numPr)
            if numPr is not None:
                num_id = numPr.find(qn('w:numId'))
                if num_id is not None and num_id.get(qn('w:val')) == '0':
                    continue 
                pPr.remove(numPr)


def insert_docx_via_xml(main_doc, target_paragraph, subdoc_bytes):
    """
    使用 docxcompose 插入子文档，自动处理样式映射、编号冲突、图片迁移。
    保留标题降级和 Body Text 2 重置逻辑。

    关键设计决策：
    - restart_numbering=False: 防止 docxcompose 重置多级列表编号，
      避免 ilvl=0 硬编码导致层级丢失（如 5.4.6.3 变成 5.5）
    - reset_reference_mapping: 每个子文档重置映射，防止 numId 冲突
    """
    sub_doc = Document(io.BytesIO(subdoc_bytes))

    target_element = target_paragraph._element
    parent_elm = target_element.getparent()

    if parent_elm is None:
        print(f"   ❌ 无法找到暗码节点的父元素")
        return 0

    insert_index = list(parent_elm).index(target_element)
    inserted_count = 0

    # 使用 docxcompose Composer 处理样式、编号、图片等
    composer = Composer(main_doc)
    # 🔑 关键：禁止编号重启，防止多级列表编号层级丢失
    # docxcompose 的 restart_first_numbering 会硬编码 ilvl=0，
    # 导致 5.4.6.3 层级丢失变成 5.5
    composer.restart_numbering = False
    # 🔑 每个子文档都需要重置映射，防止 numId 映射冲突
    composer.reset_reference_mapping()
    composer._create_style_id_mapping(sub_doc)

    for child in sub_doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'sectPr':
            continue

        if tag not in ('p', 'tbl'):
            continue

        # 深度拷贝并应用 docxcompose 转换（样式、编号、图片、形状、图表）
        new_element = copy.deepcopy(child)
        composer.add_styles(sub_doc, new_element)
        composer.add_numberings(sub_doc, new_element)
        composer.add_images(sub_doc, new_element)
        composer.add_shapes(sub_doc, new_element)
        composer.add_diagrams(sub_doc, new_element)

        # 标题降级、Body Text 2 重置（docxcompose 不处理这个）
        shift_heading_levels(new_element, main_doc)
# ========================================================
        # 👇👇👇 方案B：抓鬼代码 👇👇👇
        # ========================================================
        paragraph_text = "".join(t.text for t in new_element.iter(qn('w:t')) if t.text)
        if "如网关故障" in paragraph_text:
            print("\n🚨🚨🚨 抓到幽灵了！马上要插入的 XML 是：\n")
            print(new_element.xml)
            print("\n🚨🚨🚨 抓鬼结束\n")
        # ========================================================
        parent_elm.insert(insert_index, new_element)
        insert_index += 1
        inserted_count += 1

    parent_elm.remove(target_element)
    return inserted_count


def find_and_replace_insert_codes(main_doc, url_mapping):
    total_merged = 0

    paragraphs_to_process = []
    
    # 1. 搜索文档主体段落
    print(f"   🔍 搜索文档主体段落: {len(main_doc.paragraphs)} 个")
    for idx, para in enumerate(main_doc.paragraphs):
        para_text = para.text
        # 🔍 调试：打印所有段落的文本内容
        if idx < 10 or '[INSERT_DOC:' in para_text or 'INSERT_DOC' in para_text or '售后' in para_text or '服务手册' in para_text:
            print(f"   🔍 段落 {idx}: {repr(para_text[:150])}")
        if '[INSERT_DOC:' in para_text or 'INSERT_DOC' in para_text:
            print(f"   🔍 主体段落包含暗号关键字: {para_text[:100]}")
        codes = re.findall(r'\[INSERT_DOC:([^\]]+)\]', para_text)
        for code in codes:
            paragraphs_to_process.append((para, code))
    
    # 2. 搜索表格内的段落
    print(f"   🔍 搜索表格: {len(main_doc.tables)} 个")
    for table_idx, table in enumerate(main_doc.tables):
        print(f"   🔍 表格 {table_idx}: {len(table.rows)} 行")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    para_text = para.text
                    # 🔍 调试：打印表格中包含关键字的段落
                    if '售后' in para_text or '服务手册' in para_text or '[INSERT_DOC:' in para_text or 'INSERT_DOC' in para_text:
                        print(f"   🔍 表格[{table_idx}][{row_idx}][{cell_idx}] 段落 {para_idx}: {repr(para_text[:150])}")
                    if '[INSERT_DOC:' in para_text or 'INSERT_DOC' in para_text:
                        print(f"   🔍 表格[{table_idx}][{row_idx}][{cell_idx}] 段落 {para_idx} 包含暗号关键字: {para_text[:100]}")
                    codes = re.findall(r'\[INSERT_DOC:([^\]]+)\]', para_text)
                    for code in codes:
                        print(f"   ✅ 找到暗号: [INSERT_DOC:{code}]")
                        paragraphs_to_process.append((para, code))

    if not paragraphs_to_process:
        print("   ℹ️ 未找到任何 [INSERT_DOC:xxx] 暗码")
        return 0

    print(f"   📍 共找到 {len(paragraphs_to_process)} 个暗码引用")

    # 🔄 倒序处理，避免插入操作导致后续暗码索引偏移
    paragraphs_to_process.reverse()

    for para, code in paragraphs_to_process:
        print(f"\n   📍 处理暗号: [INSERT_DOC:{code}]")

        url = find_url_fuzzy(code, url_mapping)

        if not url:
            print(f"   ❌ 未找到对应URL，替换为错误提示文本")
            para.clear()
            para.add_run(f"[未找到附件: {code}]")
            continue

        print(f"   🔗 URL: {url}")

        try:
            resp = requests.get(url, timeout=60, verify=False)
            if len(resp.content) < 1000:
                print(f"   ⚠️ 响应内容太小，可能不是有效的docx文件")
                para.clear()
                para.add_run(f"[附件无效: {code}]")
                continue

            resp.raise_for_status()
        except Exception as e:
            print(f"   ❌ 下载失败: {e}")
            para.clear()
            para.add_run(f"[下载失败: {code}]")
            continue

        try:
            count = insert_docx_via_xml(main_doc, para, resp.content)
            print(f"   ✅ 成功插入 {count} 个节点")
            total_merged += 1
        except Exception as e:
            print(f"   ❌ XML合并失败: {e}")
            para.clear()
            para.add_run(f"[合并失败: {code}]")

    return total_merged


# ===== 动态表格行克隆引擎 =====

def get_tc_gridspan(tc_element):
    grid_span_elem = tc_element.find(qn('w:gridSpan'))
    if grid_span_elem is not None:
        val = grid_span_elem.get(qn('w:val'))
        return int(val) if val else 1
    return 1


def set_cell_text(tc_element, text):
    for p in tc_element.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                r.remove(t)

    paragraphs = tc_element.findall(qn('w:p'))
    if not paragraphs:
        new_p = OxmlElement('w:p')
        tc_element.append(new_p)
        paragraphs = [new_p]

    first_p = paragraphs[0]
    runs = first_p.findall(qn('w:r'))
    if not runs:
        new_r = OxmlElement('w:r')
        first_p.append(new_r)
        runs = [new_r]

    t_elem = OxmlElement('w:t')
    t_elem.text = str(text)
    t_elem.set(qn('xml:space'), 'preserve')
    runs[0].append(t_elem)


def get_cell_text_from_element(tc_element):
    texts = []
    for t in tc_element.iter(qn('w:t')):
        if t.text:
            texts.append(t.text)
    return ''.join(texts).strip()


def build_header_texts_with_span(header_row_element):
    header_texts = []
    for tc in header_row_element.findall(qn('w:tc')):
        text = get_cell_text_from_element(tc).strip()
        span = get_tc_gridspan(tc)
        for _ in range(span):
            header_texts.append(text)
    return header_texts


def find_template_row_index(table):
    for i in range(1, len(table.rows)):
        row = table.rows[i]
        has_empty = False
        has_content = False
        for cell in row.cells:
            text = cell.text.strip()
            if not text or re.match(r'^[\s\u3000_\-－]+$', text):
                has_empty = True
            else:
                has_content = True
        if has_empty or not has_content:
            return i
    return 1


def normalize_header(text):
    """归一化表头：去除所有空格、换行符、制表符"""
    if not text:
        return ""
    return re.sub(r'\s+', '', str(text).strip())


def parse_filled_html_to_cells(html_string):
    """
    解析Dify返回的HTML表格，提取单元格文本内容为二维数组
    
    Args:
        html_string: HTML表格字符串
        
    Returns:
        list[list[str]]: 二维数组，每个元素是单元格文本
    """
    if not html_string or not html_string.strip():
        return []
    
    try:
        soup = BeautifulSoup(html_string, 'html.parser')
        table = soup.find('table')
        
        if not table:
            print("   ⚠️ HTML中未找到table标签")
            return []
        
        rows = []
        for tr in table.find_all('tr'):
            cells = []
            for td in tr.find_all(['td', 'th']):
                # 获取单元格文本，去除首尾空白
                text = td.get_text(strip=True)
                cells.append(text)
            if cells:  # 只添加非空行
                rows.append(cells)
        
        print(f"   ✅ HTML解析成功: {len(rows)} 行")
        return rows
        
    except Exception as e:
        print(f"   ❌ HTML解析失败: {e}")
        return []


def fill_dynamic_tables(doc, dynamic_tables):
    if not dynamic_tables:
        return 0

    total_filled = 0

    for table_info in dynamic_tables:
        table_id = table_info.get("table_id")
        filled_html = table_info.get("filled_html")  # 🆕 接收HTML而不是rows
        rows_data = table_info.get("rows", [])  # 保留向后兼容
        append_images = table_info.get("append_images", [])
        append_images_by_person = table_info.get("append_images_by_person", {})
        fill_mode = table_info.get("fill_mode", "multi_person")

        print(f"   🔍 append_images: {append_images}")
        print(f"   🔍 append_images_by_person: {append_images_by_person}")

        if table_id is None or table_id >= len(doc.tables):
            print(f"   ⚠️ table_id {table_id} 超出范围，跳过")
            continue

        # 🆕 优先使用HTML，如果没有则使用rows（向后兼容）
        if filled_html:
            print(f"   📊 表格 {table_id} 使用HTML填充模式")
            rows_data = parse_filled_html_to_cells(filled_html)
            if not rows_data:
                print(f"   ⚠️ HTML解析失败，跳过表格 {table_id}")
                continue
        elif not rows_data:
            print(f"   ⚠️ 表格 {table_id} 既没有HTML也没有rows数据，跳过")
            continue

        table = doc.tables[table_id]

        # GridSpan 感知的表头提取
        header_row_element = table.rows[0]._element
        header_texts = build_header_texts_with_span(header_row_element)
        print(f"   📊 表格 {table_id} 表头(含跨列): {header_texts}")
        print(f"   🔧 填充模式: {fill_mode}")

        # 🆕 HTML模式：直接按行填充，不需要归一化映射
        if filled_html:
            print(f"   📋 HTML填充模式：直接按行列对应填充")
            
            # 从第0行开始填充（简历表所有行都是数据行，无表头行需跳过）
            for row_idx, row_cells in enumerate(rows_data):
                if row_idx >= len(table.rows):
                    print(f"   ⚠️ HTML行 {row_idx} 超出Word表格范围，跳过")
                    break
                
                target_row = table.rows[row_idx]
                word_cells = target_row._element.findall(qn('w:tc'))
                
                filled_count = 0
                for col_idx, cell_value in enumerate(row_cells):
                    if col_idx >= len(word_cells):
                        break
                    
                    # 过滤[空白]标记和空字符串
                    if cell_value and str(cell_value).strip() and str(cell_value).strip() != "[空白]":
                        tc = word_cells[col_idx]
                        set_cell_text(tc, str(cell_value))
                        filled_count += 1
                        print(f"      ✓ [{row_idx},{col_idx}] = '{cell_value}'")
                
                print(f"   ✅ 表格 {table_id} 行 {row_idx}: 填充 {filled_count}/{len(row_cells)} 个单元格")
            
            total_filled += len(rows_data)
            print(f"   ✅ 表格 {table_id} HTML填充完成: {len(rows_data)} 行")

        else:
            # 🆕 构建归一化映射表（仅用于旧的rows模式）
            normalized_map = {}
            for h in header_texts:
                norm_key = normalize_header(h)
                if norm_key:  # 只添加非空的归一化key
                    normalized_map[norm_key] = h
            print(f"   🔧 归一化映射: {normalized_map}")

            # ===== 根据填充模式选择不同的处理逻辑 =====
            if fill_mode == "multi_person":
                # ===== 汇总表模式：填充现有行，不删除模板行 =====
                print(f"   📋 汇总表模式：填充现有空白行")
                
                for i, row_data in enumerate(rows_data):
                    # 获取目标行索引（前端传来的 _rowIndex，或默认顺序填充）
                    target_row_index = row_data.get("_rowIndex", i + 1)
                    
                    if target_row_index >= len(table.rows):
                        print(f"   ⚠️ 行索引 {target_row_index} 超出范围（表格只有 {len(table.rows)} 行），跳过")
                        continue
                    
                    # 归一化前端数据的keys
                    normalized_row_data = {}
                    for key, value in row_data.items():
                        if key.startswith("_"):  # 跳过元数据字段
                            continue
                        norm_key = normalize_header(key)
                        if norm_key in normalized_map:
                            original_header = normalized_map[norm_key]
                            normalized_row_data[original_header] = value
                        else:
                            normalized_row_data[key] = value
                    
                    print(f"   🔄 第 {i+1} 行归一化: {len(row_data)} keys → {len(normalized_row_data)} keys")
                    
                    # 获取目标行的单元格
                    target_row = table.rows[target_row_index]
                    cells = target_row._element.findall(qn('w:tc'))
                    
                    # 填充单元格
                    logical_col = 0
                    filled_count = 0
                    for tc in cells:
                        span = get_tc_gridspan(tc)
                        header = header_texts[logical_col] if logical_col < len(header_texts) else ""
                        value = normalized_row_data.get(header, "")
                        
                        # 过滤[空白]标记和空字符串
                        if value and str(value).strip() and str(value).strip() != "[空白]":
                            original_text = get_cell_text_from_element(tc)
                            set_cell_text(tc, str(value))
                            filled_count += 1
                            print(f"      ✓ [{target_row_index},{logical_col}] {header}: '{original_text}' → '{value}'")
                        
                        logical_col += span
                    
                    print(f"   ✅ 表格 {table_id} 行 {target_row_index}: 填充 {filled_count}/{len(cells)} 个单元格")
                
                total_filled += len(rows_data)
                print(f"   ✅ 表格 {table_id} 汇总表填充完成: {len(rows_data)} 行")
            
            else:
                # ===== 单人简历表模式：克隆模板行（原有逻辑） =====
                print(f"   📋 单人简历表模式：克隆模板行")
                
                template_idx = find_template_row_index(table)
                template_row = table.rows[template_idx]
                template_element = template_row._element

                parent = template_element.getparent()
                template_index_in_parent = list(parent).index(template_element)

                parent.remove(template_element)

                for i, row_data in enumerate(rows_data):
                    # 归一化前端数据的keys
                    normalized_row_data = {}
                    for key, value in row_data.items():
                        if key.startswith("_"):  # 跳过元数据字段
                            continue
                        norm_key = normalize_header(key)
                        if norm_key in normalized_map:
                            original_header = normalized_map[norm_key]
                            normalized_row_data[original_header] = value
                        else:
                            normalized_row_data[key] = value
                    
                    print(f"   🔄 第 {i+1} 行归一化: {len(row_data)} keys → {len(normalized_row_data)} keys")

                    new_row = copy.deepcopy(template_element)
                    cells = new_row.findall(qn('w:tc'))

                    # 智能填充：只填充有值的单元格
                    logical_col = 0
                    filled_count = 0
                    for tc in cells:
                        span = get_tc_gridspan(tc)
                        header = header_texts[logical_col] if logical_col < len(header_texts) else ""
                        value = normalized_row_data.get(header, "")

                        # 过滤[空白]标记和空字符串
                        if value and str(value).strip() and str(value).strip() != "[空白]":
                            original_text = get_cell_text_from_element(tc)
                            set_cell_text(tc, str(value))
                            filled_count += 1
                            print(f"      ✓ [{template_idx + i},{logical_col}] {header}: '{original_text}' → '{value}'")
                        else:
                            original_text = get_cell_text_from_element(tc)
                            if original_text:
                                print(f"      ⏭️ [{template_idx + i},{logical_col}] {header}: 保留原值 '{original_text}'")

                        logical_col += span

                    parent.insert(template_index_in_parent + i, new_row)
                    print(f"   ✅ 表格 {table_id} 第 {i+1} 行: 填充 {filled_count}/{len(cells)} 个单元格")

                total_filled += len(rows_data)
                print(f"   ✅ 表格 {table_id} 克隆填充 {len(rows_data)} 行")

        # ===== 在表格下方追加图片（按人员分组） =====
        images_to_append = []

        # 优先使用按人员分组的格式
        if append_images_by_person and isinstance(append_images_by_person, dict):
            for person_name, imgs in append_images_by_person.items():
                if isinstance(imgs, list):
                    for img in imgs:
                        # 支持新格式 {url, type} 和旧格式纯字符串
                        if isinstance(img, dict):
                            images_to_append.append((person_name, img.get('url', ''), img.get('type', '')))
                        else:
                            images_to_append.append((person_name, str(img), ''))
        elif append_images:
            # 向后兼容旧的扁平数组格式
            for url in append_images:
                images_to_append.append((None, url, ''))

        # 图片尺寸规范（单位：cm）
        # 身份证：6.80×4.24  营业执照：9.55×9.55  资质证书：14.63×20.66  学历证书：14.63×10.14
        def get_image_size(att_type):
            if att_type in ('id_card_front', 'id_card_back', 'id_card'):
                return (Cm(6.80), Cm(4.24))
            if att_type == 'business_license':
                return (Cm(9.55), Cm(9.55))
            if att_type in ('qualification_certificate', 'title_certificate', 'certificate'):
                return (Cm(14.63), Cm(20.66))
            if att_type in ('degree_certificate', 'graduation_certificate'):
                return (Cm(14.63), Cm(10.14))
            # 兜底：只设宽度，高度自适应
            return (Cm(14.0), None)

        if images_to_append and table_id < len(doc.tables):
            table_element = doc.tables[table_id]._element
            current = table_element

            for idx, (person_name, img_url, att_type) in enumerate(images_to_append):
                if not img_url:
                    continue
                try:
                    resp = requests.get(img_url, verify=False, timeout=15)
                    resp.raise_for_status()
                    img_stream = io.BytesIO(resp.content)

                    p_elem = OxmlElement('w:p')
                    current.addnext(p_elem)

                    p = Paragraph(p_elem, doc._body)
                    run = p.add_run()

                    width, height = get_image_size(att_type)
                    if height:
                        run.add_picture(img_stream, width=width, height=height)
                    else:
                        run.add_picture(img_stream, width=width)

                    current = p_elem
                    label = f"{person_name} - " if person_name else ""
                    print(f"   🖼️ 表格 {table_id} 追加图片 {label}{idx + 1}/{len(images_to_append)} [{att_type}] {width}×{height}: {img_url[:60]}")

                except Exception as e:
                    print(f"   ⚠️ 下载图片失败 [{img_url[:60]}]: {e}")

    return total_filled


# ===== API 端点 =====

@router.post("/merge-docs")
async def merge_docs(
    file: UploadFile = File(...),
    mapping: str = Form(""),
    dynamic_tables: str = Form("")
):
    print("\n==========================")
    print("📥 文档合并请求")
    print("==========================")

    try:
        doc_bytes = await file.read()
        print(f"📥 主文档: {len(doc_bytes)} bytes")

        main_doc = Document(io.BytesIO(doc_bytes))
        print(f"📄 段落数: {len(main_doc.paragraphs)}, 表格数: {len(main_doc.tables)}")

        # 1. 处理动态表格
        parsed_dynamic = []
        if dynamic_tables:
            try:
                parsed_dynamic = json.loads(dynamic_tables)
                print(f"📋 动态表格数据: {len(parsed_dynamic)} 个表格")
            except json.JSONDecodeError as e:
                print(f"⚠️ dynamic_tables JSON解析失败: {e}")

        if parsed_dynamic:
            print(f"\n{'='*60}")
            print("🔧 开始填充动态表格")
            print(f"{'='*60}")
            filled = fill_dynamic_tables(main_doc, parsed_dynamic)
            print(f"✅ 动态表格填充完成，共 {filled} 行")

        # 2. 处理服务手册暗码合并
        url_mapping = {}
        if mapping:
            try:
                url_mapping = json.loads(mapping)
            except json.JSONDecodeError:
                url_mapping = {}

        if url_mapping:
            print(f"\n{'='*60}")
            print("🔍 扫描 [INSERT_DOC:xxx] 暗码并合并子文档")
            print(f"{'='*60}")
            print(f"📋 mapping: {len(url_mapping)} 条")
            total_merged = find_and_replace_insert_codes(main_doc, url_mapping)
            print(f"📦 合并完成，共 {total_merged} 个附件")

        output_io = io.BytesIO()
        main_doc.save(output_io)
        output_io.seek(0)

        final_size = len(output_io.getvalue())
        print(f"✅ 最终文档大小: {final_size} bytes\n")

        return StreamingResponse(
            output_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Merged_Bid.docx"}
        )

    except Exception as e:
        print(f"💥 服务器错误: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )
