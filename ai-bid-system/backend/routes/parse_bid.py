from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from docx import Document
from io import BytesIO
from docx.text.paragraph import Paragraph  # 👈 新增这一行：引入 Paragraph 对象

from schemas.response import (
    NormalBlank, DynamicTable, ManualTable, BlankCell, TableStructure, ParseResponse
)
from services.blank_scanner import scan_normal_blanks
from services.table_parser import parse_table, extract_table_headers
from services.table_classifier import classify_table, get_table_type_label

router = APIRouter()


@router.post("/parse-bid-docx", response_model=ParseResponse)
async def parse_bid_docx(file: UploadFile = File(...)):
    print(f"\n{'='*60}")
    print(f"📥 [后端] 收到解析请求")
    print(f"📥 [后端] 文件名: {file.filename}")
    print(f"📥 [后端] Content-Type: {file.content_type}")
    print(f"{'='*60}")

    if not file.filename.lower().endswith(".docx"):
        print(f"❌ [后端] 文件格式错误: {file.filename}")
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式文件")

    try:
        print(f"🔍 [DEBUG] 开始读取文件内容...")
        contents = await file.read()
        print(f"📥 [后端] 文件大小: {len(contents)} bytes")
        
        print(f"🔍 [DEBUG] 开始解析 docx 文档...")
        doc = Document(BytesIO(contents))
        print(f"✅ [DEBUG] docx 文档解析成功")
    except Exception as e:
        # ... (报错处理) ...
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(e)}")

    # ✅ 修改开始：强制获取 XML 中所有的 <w:p> 节点（包含表格内的段落，与前端正则行为保持绝对一致）
    all_p_elements = doc._element.xpath('//w:p')
    # 将 lxml 元素重新封装为 python-docx 的 Paragraph 对象
    paragraphs = [Paragraph(p, doc._part) for p in all_p_elements]
    # ✅ 修改结束

    print(f"📄 [后端] 真实段落总数(含表格内): {len(paragraphs)}, 表格数: {len(doc.tables)}")

    normal_blanks_raw = scan_normal_blanks(paragraphs)
    normal_blanks = [NormalBlank(**b) for b in normal_blanks_raw]
    print(f"📊 [后端] 空白扫描完成: {len(normal_blanks)} 个普通填空")

    table_structures = []
    dynamic_tables = []
    manual_tables = []

    for idx, table in enumerate(doc.tables):
        try:
            table_info = parse_table(table, idx, doc)
            
            headers = table_info["headers"]
            anchor = table_info["anchorContext"]
            # 🆕 收集所有单元格文本，用于合并单元格表格的分类
            cell_texts = [cell["text"] for cell in table_info["cells"] if cell.get("text")]
            table_type = classify_table(anchor, headers, cell_texts)
            type_label = get_table_type_label(anchor, headers)
            print(f"📊 [后端] 表格 {idx}: type={table_type}, label={type_label}, rows={table_info['rowCount']}, anchor=\"{anchor[:60]}\"")
            print(f"📊 [后端] 表格 {idx} headers: {headers}")
            print(f"📊 [后端] 表格 {idx} cell_texts样本: {cell_texts[:10]}")
            print(f"📊 [后端] 表格 {idx} blankCells: {len(table_info['blankCells'])}")

            table_structure = TableStructure(
                tableId=idx,
                rowCount=table_info["rowCount"],
                colCount=table_info["colCount"],
                headers=headers,
                cells=table_info["cells"],
                blankCells=[BlankCell(**bc) for bc in table_info["blankCells"]],
                anchorContext=anchor,
                tableHtml=table_info.get("tableHtml", ""),  # 🆕 传递完整的表格HTML
            )
            table_structures.append(table_structure)

            if table_type == "dynamic":
                dynamic_tables.append(DynamicTable(
                    tableId=idx,
                    type=type_label,
                    anchorContext=anchor,
                    headers=headers,
                    rowCount=table_info["rowCount"],
                    blankCells=[BlankCell(**bc) for bc in table_info["blankCells"]],
                    fillMode=table_info.get("fillMode", "multi_person"),  # 🆕 填充模式
                    emptyRowCount=table_info.get("emptyRowCount", 0),  # 🆕 空白行数
                    tableHtml=table_info.get("tableHtml", ""),  # 🆕 传递完整的表格HTML
                ))
                print(f"   🔧 [后端] 表格 {idx} fillMode={table_info.get('fillMode')}, emptyRows={table_info.get('emptyRowCount')}")
            else:
                manual_tables.append(ManualTable(
                    tableId=idx,
                    type=type_label,
                    anchorContext=anchor,
                    headers=headers,
                ))
        except Exception as e:
            print(f"⚠️ [后端] 解析表格 {idx} 失败，已跳过: {e}")
            continue  # 遇到无法解析的复杂表格，跳过而不是让整个接口崩溃

    meta = {
        "fileName": file.filename,
        "totalParagraphs": len(paragraphs),
        "totalTables": len(doc.tables),
        "totalNormalBlanks": len(normal_blanks),
        "totalDynamicTables": len(dynamic_tables),
        "totalManualTables": len(manual_tables),
    }

    print(f"\n{'='*60}")
    print(f"✅ [后端] 解析完成!")
    print(f"   普通填空: {meta['totalNormalBlanks']}")
    print(f"   动态表格: {meta['totalDynamicTables']}")
    print(f"   高危表格: {meta['totalManualTables']}")
    print(f"{'='*60}\n")

    return ParseResponse(
        success=True,
        normalBlanks=normal_blanks,
        dynamicTables=dynamic_tables,
        manualTables=manual_tables,
        tableStructures=table_structures,
        meta=meta,
    )
