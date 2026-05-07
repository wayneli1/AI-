"""
文本替换工具模块
提供保留格式的段落文本替换功能，支持跨run的文本替换
"""

import io
import re
import requests
from docx.shared import Cm, Inches
from docx.oxml.ns import qn


def replace_text_preserve_format(paragraph, old_text, new_text):
    """
    在段落中替换文本，保留所有格式（加粗、下划线、字体等）
    支持处理跨run的情况
    
    Args:
        paragraph: python-docx Paragraph 对象
        old_text: 要替换的原始文本
        new_text: 替换后的新文本
        
    Returns:
        bool: 是否成功替换
    """
    if not old_text:
        # 如果 old_text 为空，说明是插入操作（如冒号后直接插入）
        if paragraph.runs:
            # 在第一个 run 的末尾插入
            paragraph.runs[0].text = paragraph.runs[0].text + new_text
            return True
        else:
            # 段落没有 run，创建一个新的
            paragraph.add_run(new_text)
            return True
    
    # 获取段落的完整文本
    full_text = paragraph.text
    
    # 检查是否包含要替换的文本
    if old_text not in full_text:
        return False
    
    # 找到替换位置
    start_idx = full_text.index(old_text)
    end_idx = start_idx + len(old_text)
    
    # 遍历 runs，找到覆盖范围
    current_pos = 0
    runs_to_modify = []
    
    for run in paragraph.runs:
        run_len = len(run.text)
        run_start = current_pos
        run_end = current_pos + run_len
        
        # 判断 run 是否与替换范围重叠
        if run_end > start_idx and run_start < end_idx:
            # 计算在当前 run 内的重叠范围
            overlap_start = max(0, start_idx - run_start)
            overlap_end = min(run_len, end_idx - run_start)
            runs_to_modify.append((run, overlap_start, overlap_end, run_start))
        
        current_pos += run_len
    
    # 执行替换
    if runs_to_modify:
        # 在第一个 run 中插入新文本
        first_run, first_start, first_end, _ = runs_to_modify[0]
        first_run.text = (
            first_run.text[:first_start] + 
            new_text + 
            first_run.text[first_end:]
        )
        
        # 清空其他 run 的重叠部分
        for run, start, end, _ in runs_to_modify[1:]:
            run.text = run.text[:start] + run.text[end:]
        
        return True
    
    return False


def _is_image_url(text):
    """判断文本是否为图片URL"""
    if not text or not isinstance(text, str):
        return False
    text = text.strip()
    if not text.startswith(('http://', 'https://')):
        return False
    lower = text.lower()
    return any(lower.endswith(ext) for ext in [
        '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg',
        '.tiff', '.tif', '.ico'
    ]) or '/image' in lower or '/storage/' in lower or 'supabase' in lower


def _guess_image_type(context, blank_type):
    """根据blank上下文和类型推断图片类别"""
    if not context:
        return 'default'
    ctx = context.lower() if isinstance(context, str) else ''
    # 身份证
    if any(kw in ctx for kw in ('身份证', '身份证明', 'id_card')):
        return 'id_card'
    # 营业执照
    if any(kw in ctx for kw in ('营业执照', '工商登记', 'business_license')):
        return 'business_license'
    # 资质证书 / 职称证书
    if any(kw in ctx for kw in ('资质', '资质证书', '资格证书', '职称', 'qualification', 'title_certificate')):
        return 'qualification_certificate'
    # 学历证书 / 毕业证书
    if any(kw in ctx for kw in ('学历', '毕业证', '学位证', 'degree', 'graduation')):
        return 'degree_certificate'
    return 'default'


# 图片尺寸规范（单位：cm）
_IMAGE_SIZE_MAP = {
    'id_card':                  (Cm(6.80),  Cm(4.24)),
    'id_card_front':            (Cm(6.80),  Cm(4.24)),
    'id_card_back':             (Cm(6.80),  Cm(4.24)),
    'business_license':         (Cm(14.65),  Cm(9.55)),
    'qualification_certificate': (Cm(14), Cm(11)),
    'title_certificate':        (Cm(14), Cm(11)),
    'certificate':               (Cm(14), Cm(11)),
    'degree_certificate':       (Cm(14.63), Cm(10.14)),
    'graduation_certificate':   (Cm(14.63), Cm(10.14)),
    'default':                  (Cm(14.0),  None),
}


def _get_image_size_for_blank(blank_type, context=''):
    """
    根据blank类型和上下文返回合适的图片尺寸（单位：cm）
    
    尺寸规范：
    - 身份证：6.80×4.24
    - 营业执照：9.55×9.55
    - 资质证书：14.63×20.66
    - 学历证书：14.63×10.14
    """
    # 先通过上下文推断图片类别
    guessed_type = _guess_image_type(context, blank_type)
    return _IMAGE_SIZE_MAP.get(guessed_type, _IMAGE_SIZE_MAP['default'])


def insert_image_at_paragraph(paragraph, image_url, blank_type='', context='', max_retries=1):
    """
    在段落中插入图片，替换原有的占位文本。
    
    Args:
        paragraph: python-docx Paragraph 对象
        image_url: 图片URL
        blank_type: blank类型（用于确定图片尺寸）
        context: blank上下文（用于推断图片类型和尺寸）
        max_retries: 下载重试次数
    
    Returns:
        bool: 是否成功插入
    """
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    for attempt in range(max_retries + 1):
        try:
            resp = requests.get(image_url, verify=False, timeout=30)
            if resp.status_code != 200:
                print(f"   ⚠️ 下载图片失败 (HTTP {resp.status_code}): {image_url[:80]}")
                continue
            img_stream = io.BytesIO(resp.content)
            
            # 验证是否为有效图片
            if len(resp.content) < 100:
                print(f"   ⚠️ 图片数据太小 ({len(resp.content)} bytes), 跳过: {image_url[:80]}")
                continue
            
            # 清空段落现有内容（保留段落属性）
            for run in paragraph.runs:
                run.text = ''
            # 清除段落中所有run元素
            p_elem = paragraph._element
            for child in list(p_elem):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'r':  # w:r = run
                    p_elem.remove(child)
            
            # 插入图片
            width, height = _get_image_size_for_blank(blank_type, context)
            run = paragraph.add_run()
            if height:
                run.add_picture(img_stream, width=width, height=height)
            else:
                run.add_picture(img_stream, width=width)
            
            print(f"   🖼️ 成功插入图片 [{blank_type}] {width}×{height}: {image_url[:80]}")
            return True
            
        except Exception as e:
            if attempt < max_retries:
                print(f"   ⚠️ 插入图片失败 (尝试 {attempt+1}): {e}")
                continue
            print(f"   ❌ 插入图片最终失败: {e}")
            return False
    
    return False


def replace_text_fuzzy(paragraph, old_text, new_text, blank_type='', context=''):
    """
    模糊替换：尝试多种策略定位并替换文本
    支持：纯文本替换 + 图片URL自动识别插入
    
    策略优先级：
    0. 如果 new_text 是图片URL → 插入图片
    1. 精确匹配
    2. 去除空格后匹配
    3. 下划线/横线模糊匹配
    
    Args:
        paragraph: python-docx Paragraph 对象
        old_text: 要替换的原始文本
        new_text: 替换后的新文本（文本或图片URL）
        blank_type: blank类型（用于确定图片尺寸）
        context: blank上下文（用于推断图片类型和尺寸）
        
    Returns:
        bool: 是否成功替换
    """
    # 策略0: 如果 new_text 是图片URL → 下载并插入图片
    if _is_image_url(new_text):
        # 同一个blank可能填了多个URL（换行分隔），逐个插入
        urls = [u.strip() for u in re.split(r'[\n\r]+', new_text) if u.strip() and u.strip().startswith(('http://', 'https://'))]
        
        if not urls:
            # URL解析失败，回退到文本替换
            pass
        elif len(urls) == 1:
            # 单个图片URL
            if old_text:
                # 先清除占位文本
                if not replace_text_preserve_format(paragraph, old_text, ''):
                    # 清除失败也没关系，继续插入图片
                    pass
            return insert_image_at_paragraph(paragraph, urls[0], blank_type, context)
        else:
            # 多个图片URL：第一张替换占位文本，后续追加新段落
            if old_text:
                replace_text_preserve_format(paragraph, old_text, '')
            
            success = insert_image_at_paragraph(paragraph, urls[0], blank_type, context)
            
            # 在当前段落后追加后续图片
            parent = paragraph._element.getparent()
            if parent is not None:
                insert_after = paragraph._element
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph
                
                for url in urls[1:]:
                    new_p = OxmlElement('w:p')
                    insert_after.addnext(new_p)
                    new_paragraph = Paragraph(new_p, paragraph._part)
                    insert_image_at_paragraph(new_paragraph, url, blank_type, context)
                    insert_after = new_p
            
            return success
    
    # 策略1: 精确匹配
    if replace_text_preserve_format(paragraph, old_text, new_text):
        return True
    
    # 策略2: 去除空格后匹配
    full_text = paragraph.text
    normalized_old = old_text.replace(' ', '').replace('\t', '').replace('\n', '')
    normalized_full = full_text.replace(' ', '').replace('\t', '').replace('\n', '')
    
    if normalized_old in normalized_full:
        # 找到归一化后的位置，映射回原始文本
        norm_idx = normalized_full.index(normalized_old)
        
        # 映射回原始索引
        original_idx = 0
        norm_count = 0
        for i, char in enumerate(full_text):
            if char not in (' ', '\t', '\n'):
                if norm_count == norm_idx:
                    original_idx = i
                    break
                norm_count += 1
        
        # 计算原始文本中对应的长度
        original_end = original_idx
        norm_matched = 0
        for i in range(original_idx, len(full_text)):
            if full_text[i] not in (' ', '\t', '\n'):
                norm_matched += 1
                if norm_matched == len(normalized_old):
                    original_end = i + 1
                    break
        
        # 提取原始匹配文本并替换
        actual_old_text = full_text[original_idx:original_end]
        if replace_text_preserve_format(paragraph, actual_old_text, new_text):
            return True
    
    # 策略3: 下划线/横线模糊匹配
    if re.match(r'^_{2,}$', old_text):
        # 匹配任意长度的下划线
        pattern = r'_{2,}'
        match = re.search(pattern, full_text)
        if match:
            matched_text = match.group()
            return replace_text_preserve_format(paragraph, matched_text, new_text)
    
    if re.match(r'^-{3,}$', old_text):
        # 匹配任意长度的横线
        pattern = r'-{3,}'
        match = re.search(pattern, full_text)
        if match:
            matched_text = match.group()
            return replace_text_preserve_format(paragraph, matched_text, new_text)
    
    return False
