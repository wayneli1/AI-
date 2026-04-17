"""
测试脚本：验证表格解析器的修复效果
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from services.table_parser import parse_table

def test_parse_document(docx_path):
    """测试解析Word文档中的表格"""
    print(f"[文档] 正在解析文档: {docx_path}")
    print("=" * 60)
    
    try:
        doc = Document(docx_path)
        print(f"[成功] 文档加载成功，共 {len(doc.tables)} 个表格\n")
        
        # 只测试前2个表格
        for idx, table in enumerate(doc.tables[:2]):
            print(f"\n{'='*60}")
            print(f"[表格] 表格 {idx}")
            print(f"{'='*60}")
            
            result = parse_table(table, idx, doc)
            
            print(f"[锚点] 锚点上下文: {result['anchorContext']}")
            print(f"[尺寸] 行数: {result['rowCount']}, 列数: {result['colCount']}")
            print(f"[表头] 表头 ({len(result['headers'])} 列):")
            for i, h in enumerate(result['headers']):
                print(f"   列{i}: '{h}'")
            
            print(f"\n[空白] 空白单元格 ({len(result['blankCells'])} 个):")
            for bc in result['blankCells'][:10]:  # 只显示前10个
                print(f"   行{bc['row']} 列{bc['col']}: label='{bc['label']}', headerText='{bc['headerText']}', rowHeader='{bc['rowHeader']}'")
            
            if len(result['blankCells']) > 10:
                print(f"   ... 还有 {len(result['blankCells']) - 10} 个空白单元格")
        
        print(f"\n{'='*60}")
        print("[完成] 测试完成！")
        print(f"{'='*60}")
        
    except Exception as e:
        print(f"[错误] 错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # 使用你的测试文档
    test_parse_document("简历模版.docx")
