# merge_server.py - 使用 docxcompose 进行文档合并
from docx.oxml.ns import qn
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docxcompose.composer import Composer
import json
import requests
import re
import io
import copy
import urllib3


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

def normalize_key(key):
    if not key:
        return key
    return re.sub(r'\s+', '', key)

def find_url_fuzzy(code, url_mapping):
    if code in url_mapping:
        return url_mapping[code]

    normalized_code = normalize_key(code)
    for key, url in url_mapping.items():
        if normalize_key(key) == normalized_code:
            return url

    full_code = f"[INSERT_DOC:{code}]"
    if full_code in url_mapping:
        return url_mapping[full_code]

    normalized_full = normalize_key(full_code)
    for key, url in url_mapping.items():
        if normalize_key(key) == normalized_full:
            return url

    return None


def _build_heading_level_map(doc):
    """
    根据文档的 styles 表动态构建样式 ID -> 标题级别的映射。
    """
    mapping = {}
    try:
        for style in doc.styles:
            if style.name:
                m = re.match(r'^Heading\s*(\d+)$', style.name, re.IGNORECASE)
                if m:
                    mapping[style.style_id] = int(m.group(1))
    except Exception as e:
        print(f"   [WARN] 构建标题映射失败: {e}")
    return mapping


def shift_heading_levels(element, doc):
    """
    合并子文档时处理段落样式：
    1. 标题降一级（Heading1→Heading2 等）
    2. 剥离标题段落的 numPr，防止编号冲突
    3. 处理 Body Text 2 样式：重置为 Normal
    """
    heading_level_map = _build_heading_level_map(doc)

    qn_pStyle = qn('w:pStyle')
    qn_numPr = qn('w:numPr')

    for p in element.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue

        style_elem = pPr.find(qn_pStyle)
        is_heading = False
        is_body_text_2 = False

        if style_elem is not None:
            val = style_elem.get(qn('w:val'), '')

            # 匹配文本格式：Heading1 / Heading 1 / Heading11 等
            m = re.match(r'^Heading\s*(\d+)$', val, re.IGNORECASE)
            if m:
                is_heading = True
                old_level = int(m.group(1))
                new_level = min(old_level + 1, 9)
                style_elem.set(qn('w:val'), f'Heading{new_level}')
            # 匹配动态构建的样式 ID 映射（docxcompose 映射后的数字 ID）
            elif val in heading_level_map:
                is_heading = True
                old_level = heading_level_map[val]
                new_level = min(old_level + 1, 9)
                new_style_id = None
                for sid, level in heading_level_map.items():
                    if level == new_level:
                        new_style_id = sid
                        break
                if new_style_id:
                    style_elem.set(qn('w:val'), new_style_id)
                else:
                    style_elem.set(qn('w:val'), f'Heading{new_level}')
            # 处理 Body Text 2 样式：重置为 Normal
            # Word 中样式 ID 可能是 "BodyText2" 或 "Body Text 2"
            elif val == '2' or val.lower().replace(' ', '') == 'bodytext2':
                is_body_text_2 = True
                style_elem.set(qn('w:val'), 'Normal')

        # 剥离标题段落的 numPr
        if is_heading or is_body_text_2:
            numPr = pPr.find(qn_numPr)
            if numPr is not None:
                pPr.remove(numPr)


def insert_docx_via_xml(main_doc, target_paragraph, subdoc_bytes):
    """
    使用 docxcompose 插入子文档，自动处理样式映射、编号冲突、图片迁移。
    保留标题降级和 Body Text 2 重置逻辑。

    关键设计决策：
    - restart_numbering=False: 防止 docxcompose 重置多级列表编号，
      避免 ilvl=0 硬编码导致层级丢失（如 5.4.6.3 变成 5.5）
    - reset_reference_mapping: 每个子文档重置映射，防止 numId 冲突
    """
    sub_doc = Document(io.BytesIO(subdoc_bytes))

    target_element = target_paragraph._element
    parent_elm = target_element.getparent()

    if parent_elm is None:
        print(f"   ❌ 无法找到暗码节点的父元素")
        return 0

    insert_index = list(parent_elm).index(target_element)
    inserted_count = 0

    # 使用 docxcompose Composer 处理样式、编号、图片等
    composer = Composer(main_doc)
    # 🔑 关键：禁止编号重启，防止多级列表编号层级丢失
    composer.restart_numbering = False
    # 🔑 每个子文档都需要重置映射，防止 numId 映射冲突
    composer.reset_reference_mapping()
    composer._create_style_id_mapping(sub_doc)

    for child in sub_doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'sectPr':
            continue

        if tag not in ('p', 'tbl'):
            continue

        # 深度拷贝并应用 docxcompose 转换
        new_element = copy.deepcopy(child)
        composer.add_styles(sub_doc, new_element)
        composer.add_numberings(sub_doc, new_element)
        composer.add_images(sub_doc, new_element)
        composer.add_shapes(sub_doc, new_element)
        composer.add_diagrams(sub_doc, new_element)

        # 标题降级、Body Text 2 重置
        shift_heading_levels(new_element, main_doc)

        parent_elm.insert(insert_index, new_element)
        insert_index += 1
        inserted_count += 1

    parent_elm.remove(target_element)
    return inserted_count


def find_and_replace_insert_codes(main_doc, url_mapping):
    """
    遍历主文档所有段落，查找 [INSERT_DOC:xxx] 暗码，
    下载对应子文档并调用 insert_docx_via_xml 进行合并。
    """
    total_merged = 0

    paragraphs_to_process = []
    for para in main_doc.paragraphs:
        para_text = para.text
        codes = re.findall(r'\[INSERT_DOC:([^\]]+)\]', para_text)
        for code in codes:
            paragraphs_to_process.append((para, code))

    if not paragraphs_to_process:
        print("   ℹ️ 未找到任何 [INSERT_DOC:xxx] 暗码")
        return 0

    print(f"   📍 共找到 {len(paragraphs_to_process)} 个暗码引用")

    # 🔄 倒序处理，避免插入操作导致后续暗码索引偏移
    paragraphs_to_process.reverse()

    for para, code in paragraphs_to_process:
        print(f"\n   📍 处理暗号: [INSERT_DOC:{code}]")

        url = find_url_fuzzy(code, url_mapping)

        if not url:
            print(f"   ❌ 未找到对应URL，替换为错误提示文本")
            para.clear()
            para.add_run(f"[未找到附件: {code}]")
            continue

        print(f"   🔗 URL: {url}")

        try:
            resp = requests.get(url, timeout=60, verify=False)
            print(f"   📊 HTTP状态码: {resp.status_code}")
            print(f"   📊 响应大小: {len(resp.content)} bytes")

            if len(resp.content) < 1000:
                print(f"   ⚠️ 响应内容太小，可能不是有效的docx文件")
                print(f"   ⚠️ 内容预览: {resp.content[:200]}")
                para.clear()
                para.add_run(f"[附件无效: {code}]")
                continue

            resp.raise_for_status()
        except Exception as e:
            print(f"   ❌ 下载失败: {e}")
            para.clear()
            para.add_run(f"[下载失败: {code}]")
            continue

        try:
            count = insert_docx_via_xml(main_doc, para, resp.content)
            print(f"   ✅ 成功插入 {count} 个节点")
            total_merged += 1
        except Exception as e:
            print(f"   ❌ XML合并失败: {e}")
            import traceback
            traceback.print_exc()
            para.clear()
            para.add_run(f"[合并失败: {code}]")

    return total_merged


@app.post("/api/merge-docs")
async def merge_docs(file: UploadFile = File(...), mapping: str = Form(...)):
    print("\n==========================")
    print("📥 成功接收到前端请求！")
    print("==========================")
    print("\n" + "=" * 70)
    print("🚀 [merge_docs] 开始处理文档合并请求（纯XML节点深度拷贝）")
    print("=" * 70)

    try:
        doc_bytes = await file.read()
        print(f"📥 接收到主文档: {len(doc_bytes)} bytes")

        url_mapping = json.loads(mapping)
        print(f"\n📋 mapping 接收结果:")
        print(f"   条目数量: {len(url_mapping)}")
        for key, url in url_mapping.items():
            print(f"   - 键: [{key}] -> {url}")

        main_doc = Document(io.BytesIO(doc_bytes))
        print(f"\n📄 主文档段落数: {len(main_doc.paragraphs)}")

        print(f"\n{'='*60}")
        print("🔍 开始扫描文档中的 [INSERT_DOC:xxx] 暗码并合并子文档")
        print(f"{'='*60}")

        total_merged = find_and_replace_insert_codes(main_doc, url_mapping)

        print(f"\n{'='*60}")
        print(f"📦 合并完成，共成功合并 {total_merged} 个附件")
        print(f"{'='*60}")

        output_io = io.BytesIO()
        main_doc.save(output_io)
        output_io.seek(0)

        final_size = len(output_io.getvalue())
        print(f"\n✅ 最终文档大小: {final_size} bytes")
        print("=" * 70 + "\n")

        return StreamingResponse(
            output_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Merged_Bid.docx"}
        )

    except Exception as e:
        print(f"\n💥 服务器致命错误: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("merge_server:app", host="0.0.0.0", port=8003, reload=True)
